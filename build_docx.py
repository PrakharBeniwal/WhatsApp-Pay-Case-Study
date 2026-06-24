# -*- coding: utf-8 -*-
"""
Markdown -> styled .docx converter (python-docx).
Usage: python build_docx.py <input.md> <output.docx>
"""
import sys, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x12, 0x5C, 0x42)

def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def add_inline(paragraph, text):
    pattern = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith('**'):
            r = paragraph.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith('`'):
            r = paragraph.add_run(tok[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(9.5)
        elif tok.startswith('*'):
            r = paragraph.add_run(tok[1:-1]); r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])

def is_table_sep(line):
    return bool(re.match(r'^\s*\|?\s*:?-{2,}.*$', line)) and set(line.strip()) <= set('|:- ')

def main():
    src, out = sys.argv[1], sys.argv[2]
    with open(src, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')
    doc = Document()
    style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(10.5)
    i, n = 0, len(lines)
    while i < n:
        line = lines[i].rstrip('\n'); stripped = line.strip()
        if not stripped:
            i += 1; continue
        if stripped == '[[PAGEBREAK]]':
            doc.add_page_break(); i += 1; continue
        if re.match(r'^\s*([-*_])\1{2,}\s*$', stripped):
            p = doc.add_paragraph(); pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'BBBBBB')
            pbdr.append(bottom); pPr.append(pbdr); i += 1; continue
        m = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if m:
            level = len(m.group(1)); text = m.group(2).strip()
            if level == 1:
                h = doc.add_heading('', level=0); r = h.add_run(text); r.font.color.rgb = ACCENT
            else:
                h = doc.add_heading('', level=min(level-1, 4)); r = h.add_run(text)
                if level == 2:
                    r.font.color.rgb = ACCENT
            i += 1; continue
        if '|' in line and i + 1 < n and is_table_sep(lines[i+1]):
            header = [c.strip() for c in stripped.strip('|').split('|')]
            i += 2; rows = []
            while i < n and '|' in lines[i] and lines[i].strip():
                rows.append([c.strip() for c in lines[i].strip().strip('|').split('|')]); i += 1
            ncol = len(header)
            table = doc.add_table(rows=1, cols=ncol)
            table.style = 'Light Grid Accent 1'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            for j, htext in enumerate(header):
                if j < ncol:
                    hdr[j].text = ''; p = hdr[j].paragraphs[0]; add_inline(p, htext)
                    for rr in p.runs:
                        rr.bold = True; rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    shade_cell(hdr[j], '125C42')
            for row in rows:
                cells = table.add_row().cells
                for j in range(ncol):
                    val = row[j] if j < len(row) else ''
                    cells[j].text = ''; add_inline(cells[j].paragraphs[0], val)
                    for rr in cells[j].paragraphs[0].runs:
                        rr.font.size = Pt(9)
            doc.add_paragraph(); continue
        if stripped.startswith('>'):
            qtext = re.sub(r'^>\s?', '', stripped)
            p = doc.add_paragraph(style='Intense Quote'); add_inline(p, qtext); i += 1; continue
        mb = re.match(r'^(\s*)[-*+]\s+(.*)$', line)
        if mb:
            indent = len(mb.group(1)); p = doc.add_paragraph(style='List Bullet')
            if indent >= 2:
                p.paragraph_format.left_indent = Inches(0.25 * (indent // 2 + 1))
            add_inline(p, mb.group(2).strip()); i += 1; continue
        mn = re.match(r'^(\s*)\d+\.\s+(.*)$', line)
        if mn:
            p = doc.add_paragraph(style='List Number'); add_inline(p, mn.group(2).strip()); i += 1; continue
        p = doc.add_paragraph(); add_inline(p, stripped); i += 1
    doc.save(out); print('Saved', out)

if __name__ == '__main__':
    main()
